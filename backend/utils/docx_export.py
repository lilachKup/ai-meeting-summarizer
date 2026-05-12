import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _is_rtl(text: str) -> bool:
    """Returns True if the text contains mostly Hebrew/Arabic characters."""
    if not text:
        return False
    rtl_chars = len(re.findall(r'[\u0590-\u05FF\u0600-\u06FF]', text))
    return rtl_chars > len(text) * 0.3


def _set_rtl_paragraph(paragraph) -> None:
    """Sets a paragraph to RTL direction in the Word document."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl_elem = OxmlElement('w:rtl')
        rPr.append(rtl_elem)


def _add_paragraph(doc, text: str, style: str = None, rtl: bool = False):
    """Adds a paragraph with optional style and RTL support."""
    p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    if rtl:
        _set_rtl_paragraph(p)
    return p


def generate_docx(transcript: str, summary: dict) -> bytes:
    rtl = _is_rtl(summary.get("summary", "") or transcript)

    doc = Document()
    doc.add_heading("Meeting Summary", 0)

    doc.add_heading("Summary", level=1)
    _add_paragraph(doc, summary.get("summary", ""), rtl=rtl)

    doc.add_heading("Participants", level=1)
    participants = summary.get("participants", [])
    if participants:
        for p in participants:
            _add_paragraph(doc, p, style="List Bullet", rtl=rtl)
    else:
        _add_paragraph(doc, "לא זוהו משתתפים" if rtl else "No participants identified", rtl=rtl)

    doc.add_heading("Decisions", level=1)
    decisions = summary.get("decisions", [])
    if decisions:
        for d in decisions:
            _add_paragraph(doc, d, style="List Bullet", rtl=rtl)
    else:
        _add_paragraph(doc, "לא זוהו החלטות" if rtl else "No decisions identified", rtl=rtl)

    doc.add_heading("Action Items", level=1)
    action_items = summary.get("action_items", [])
    if action_items:
        for item in action_items:
            if isinstance(item, dict):
                owner = item.get('owner', 'Not specified')
                deadline = item.get('deadline', 'Not specified')
                text = f"{item.get('task', '')} — {owner} | {deadline}"
            else:
                text = str(item)
            _add_paragraph(doc, text, style="List Bullet", rtl=rtl)
    else:
        _add_paragraph(doc, "לא זוהו משימות" if rtl else "No action items identified", rtl=rtl)

    doc.add_heading("Full Transcript", level=1)
    transcript_rtl = _is_rtl(transcript)
    _add_paragraph(doc, transcript, rtl=transcript_rtl)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()